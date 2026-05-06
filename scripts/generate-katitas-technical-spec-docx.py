#!/usr/bin/env python3
"""
Japanoma — Katitas Technical Spec (editable Word version).

Mirrors docs/Japanoma-Katitas-Technical-Spec.pdf content so Kaz can edit
freely before sending to Katitas.

Output: docs/Japanoma-Katitas-Technical-Spec.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = PROJECT_ROOT / "docs" / "Japanoma-Katitas-Technical-Spec.docx"

SUMI = RGBColor(0x1A, 0x18, 0x16)
SUMI_LIGHT = RGBColor(0x3D, 0x38, 0x33)
STONE = RGBColor(0x8A, 0x82, 0x79)
SUGI = RGBColor(0x9B, 0x7B, 0x4F)
SHOJI_HEX = "FAFAF7"


def shade(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_margins(doc: Document, cm_val: float = 2.0):
    for section in doc.sections:
        section.top_margin = Cm(cm_val)
        section.bottom_margin = Cm(cm_val)
        section.left_margin = Cm(cm_val)
        section.right_margin = Cm(cm_val)


def base_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = SUMI_LIGHT

    for name, size, color, bold, italic in [
        ("H1Title", 24, SUMI, False, False),
        ("H2Sub", 12, STONE, False, True),
        ("Section", 8, SUGI, True, False),
        ("Heading3Custom", 10, SUMI, True, False),
    ]:
        if name in doc.styles:
            continue
        style = doc.styles.add_style(name, 1)
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = bold
        style.font.italic = italic


def add_section_label(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = SUGI
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_h1(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.font.name = "Cambria"
    run.font.color.rgb = SUMI
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_sub(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = STONE
    run.font.name = "Cambria"
    p.paragraph_format.space_after = Pt(12)
    return p


def add_body(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = SUMI_LIGHT
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    # Reset run color
    if p.runs:
        p.runs[0].clear()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = SUMI_LIGHT
    return p


def add_bullets(doc, items):
    for it in items:
        add_bullet(doc, it)


def add_code_block(doc, text: str, label=None):
    if label:
        p = doc.add_paragraph()
        run = p.add_run(label.upper())
        run.font.size = Pt(7)
        run.font.color.rgb = STONE
        run.font.bold = True
        p.paragraph_format.space_after = Pt(2)

    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, SHOJI_HEX)
    cell.text = ""
    for i, line in enumerate(text.splitlines()):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = SUMI
        p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_callout(doc, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, SHOJI_HEX)
    cell.text = ""
    p_title = cell.paragraphs[0]
    run = p_title.add_run(title.upper())
    run.font.size = Pt(8)
    run.font.color.rgb = SUGI
    run.font.bold = True
    p_body = cell.add_paragraph()
    run = p_body.add_run(body)
    run.font.size = Pt(10)
    run.font.color.rgb = SUMI
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_fact_row(doc, items):
    table = doc.add_table(rows=2, cols=len(items))
    table.autofit = True
    for i, (label, value) in enumerate(items):
        lcell = table.cell(0, i)
        vcell = table.cell(1, i)
        lp = lcell.paragraphs[0]
        lrun = lp.add_run(label.upper())
        lrun.font.size = Pt(7)
        lrun.font.color.rgb = STONE
        lrun.font.bold = True
        vp = vcell.paragraphs[0]
        vrun = vp.add_run(value)
        vrun.font.size = Pt(12)
        vrun.font.name = "Cambria"
        vrun.font.color.rgb = SUMI
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_schema_table(doc, header, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = SUMI
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            if c == 0:
                run.font.name = "Consolas"
                run.font.color.rgb = SUMI
            elif c == 1:
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                run.font.color.rgb = STONE
            elif c == 2:
                run.font.bold = True
                run.font.color.rgb = SUMI
            else:
                run.font.color.rgb = SUMI_LIGHT
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build(doc: Document):
    # Header for every page
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    run = hp.add_run("JAPANOMA  ·  KATITAS TECHNICAL SPEC")
    run.font.size = Pt(8)
    run.font.color.rgb = SUGI
    run.font.bold = True
    hp.add_run("\t\tDRAFT V2  ·  APRIL 2026").font.size = Pt(8)

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fr = fp.add_run("For Katitas engineering  ·  via Go&C Partners")
    fr.font.size = Pt(8)
    fr.font.color.rgb = STONE

    # --- Page 1
    add_section_label(doc, "Katitas Technical Spec · 01")
    add_h1(doc, "Japanoma ↔ Katitas integration spec.")
    add_sub(doc, "Reference specification for engineering. Field tables, sync "
                 "protocol, URL format, and auth. No commercial context.")
    add_fact_row(doc, [
        ("Model", "Read-only referral\n(one-way, pull)"),
        ("Direction", "Katitas → Japanoma\n(listing feed)"),
        ("Refresh", "Daily delta\n+ weekly full"),
    ])
    add_fact_row(doc, [
        ("Regions", "Nagano, Niigata,\nHokkaido"),
        ("Volume", "Low thousands\nof active listings"),
        ("User PII flow", "None\n(opaque IDs only)"),
    ])
    add_section_label(doc, "Scope")
    add_body(doc, "Japanoma pulls a filtered subset of active Katitas listings, "
                  "indexes them locally, runs a matching algorithm against "
                  "signed-in user profiles, and surfaces 3–5 best-fit matches "
                  "per user. Each match deep-links back to the canonical Katitas "
                  "listing URL. No user data is sent to Katitas; no Katitas "
                  "customer data is pulled.")
    add_section_label(doc, "Out of scope")
    add_bullets(doc, [
        "Full listing re-display, including full photo sets and full descriptions.",
        "Image caching or re-hosting on Japanoma's CDN.",
        "Redistribution of the feed to third parties.",
        "Any user PII transfer in either direction.",
        "Real-time or sub-hourly sync.",
        "Writes back into the Katitas system.",
    ])
    add_section_label(doc, "What Katitas must provide")
    add_bullets(doc, [
        "One of the three data access methods on page 2.",
        "The full listing schema on pages 3–4, or the minimum viable subset.",
        "A stable listing_id and a reliable updated_at timestamp.",
        "A sandbox or sample payload for build-phase testing.",
        "A named engineering contact for the implementation window and ongoing support.",
    ])
    add_section_label(doc, "What Japanoma commits to")
    add_bullets(doc, [
        "A single technical point of contact (Craefto) for the duration of the integration.",
        "No more than the specified sync rate against the Katitas endpoint.",
        "Same-day takedown of any listing Katitas requests we remove.",
        "Attribution (\"via Katitas\") on every match card that uses Katitas data.",
        "No storage of any listing fields beyond those on pages 3–4.",
    ])
    doc.add_page_break()

    # --- Page 2
    add_section_label(doc, "Katitas Technical Spec · 02")
    add_h1(doc, "Data access.")
    add_section_label(doc, "Option A · Authenticated REST API (preferred)")
    add_body(doc, "Read-only HTTPS endpoint returning JSON. Japanoma "
                  "authenticates with a shared API key in the Authorization "
                  "header. Supports pagination and delta queries.")
    add_code_block(doc,
        "GET https://api.katitas.jp/v1/listings\n"
        "    ?region=nagano,niigata,hokkaido\n"
        "    &status=active\n"
        "    &updated_since=2026-04-09T00:00:00Z\n"
        "    &page=1\n"
        "    &per_page=100\n"
        "\n"
        "Authorization: Bearer <API_KEY>\n"
        "Accept: application/json\n"
        "\n"
        "Response 200 OK:\n"
        "{\n"
        "    \"data\": [ { ...listing object... }, ... ],\n"
        "    \"pagination\": {\n"
        "        \"page\": 1,\n"
        "        \"per_page\": 100,\n"
        "        \"total\": 847,\n"
        "        \"next_page\": 2\n"
        "    }\n"
        "}",
        label="Example delta query (illustrative — final URL TBD by Katitas)")
    add_section_label(doc, "Option B · Structured feed via SFTP or S3")
    add_body(doc, "Katitas drops a full daily snapshot (JSON or XML) to an "
                  "agreed SFTP server or private S3 bucket. Japanoma pulls on a "
                  "cron, diffs against the previous snapshot, and updates the "
                  "local index. Same schema as Option A.")
    add_code_block(doc,
        "Path:         sftp://feeds.katitas.jp/japanoma/YYYY-MM-DD/listings.json\n"
        "Auth:         SSH key, Japanoma public key added to ~/.ssh/authorized_keys\n"
        "File format:  UTF-8 JSON, one array of listing objects\n"
        "Size est:     < 10 MB uncompressed for the full snapshot",
        label="Feed delivery")
    add_section_label(doc, "Option C · Weekly CSV by email (bridge only)")
    add_body(doc, "Katitas emails a weekly CSV to a designated Japanoma "
                  "address. Japanoma manually ingests. Acceptable only as a "
                  "90-day bridge while Option A or B is being built. Same field "
                  "set as pages 3–4, one row per listing, UTF-8 with a header row.")
    add_callout(doc, "Preference order",
        "A > B > C. Option A is chosen whenever Katitas can support it. "
        "Option B is operationally simpler but puts delta detection on "
        "Japanoma's side. Option C is a fallback only.")
    doc.add_page_break()

    # --- Page 3: Schema identity/location + property
    add_section_label(doc, "Katitas Technical Spec · 03")
    add_h1(doc, "Listing schema.")
    add_section_label(doc, "Identity · Location")
    add_schema_table(doc,
        ["Field", "Type", "Req.", "Description"],
        [
            ["listing_id", "string", "Req",
             "Stable, unique, permanent identifier. Primary key in our index. Must not change across updates or re-listings."],
            ["canonical_url", "string (URL)", "Req",
             "HTTPS URL of the Katitas listing page. Deep-link target for user jump-off."],
            ["prefecture", "string", "Req",
             "Prefecture name (Japanese or romaji) or ISO-like code. Used for region filtering."],
            ["city", "string", "Req",
             "City name (Japanese or romaji). Drives city-level matching against quiz recommendations."],
            ["address_line", "string", "Opt",
             "Full street address. Not user-facing; used only for disambiguation."],
            ["lat", "number", "Pref",
             "Latitude in decimal degrees. Enables distance-to-ski-lift scoring."],
            ["lng", "number", "Pref",
             "Longitude in decimal degrees."],
        ])
    add_section_label(doc, "Property attributes")
    add_schema_table(doc,
        ["Field", "Type", "Req.", "Description"],
        [
            ["property_type", "enum (string)", "Req",
             "One of: detached_house | apartment | land | mixed_use | other. Katitas's native taxonomy is also acceptable — Japanoma will map it."],
            ["condition", "enum (string)", "Req",
             "One of: as_is | livable | light_reno | major_reno | new_build. Katitas's native taxonomy acceptable."],
            ["year_built", "integer", "Pref", "Four-digit year. Null if unknown."],
            ["land_area_sqm", "number", "Pref", "Land area in square metres. Null if not applicable."],
            ["building_area_sqm", "number", "Pref", "Total building floor area in square metres."],
            ["layout", "string", "Opt", "Room layout code, e.g. \"3LDK\", \"4LDK\"."],
        ])
    doc.add_page_break()

    # --- Page 4: Schema commercial/media/status + example
    add_section_label(doc, "Katitas Technical Spec · 04")
    add_h1(doc, "Listing schema, continued.")
    add_section_label(doc, "Commercial")
    add_schema_table(doc,
        ["Field", "Type", "Req.", "Description"],
        [
            ["asking_price_jpy", "integer", "Req",
             "Current asking price in Japanese yen (not formatted, no commas)."],
            ["price_currency", "string", "Req",
             "Always \"JPY\" for Katitas. Explicit for schema portability."],
            ["listing_agent_name", "string", "Pref",
             "Katitas office or agent name. Internal use only, not user-facing."],
        ])
    add_section_label(doc, "Media")
    add_schema_table(doc,
        ["Field", "Type", "Req.", "Description"],
        [
            ["hero_image_url", "string (URL)", "Req",
             "Single primary image URL on a Katitas-controlled domain. Hot-linked on match cards with \"via Katitas\" attribution. Not cached or re-hosted."],
            ["image_alt_text", "string", "Opt",
             "Alt text for accessibility. Falls back to listing title if absent."],
        ])
    add_section_label(doc, "Status · Timestamps")
    add_schema_table(doc,
        ["Field", "Type", "Req.", "Description"],
        [
            ["status", "enum (string)", "Req",
             "One of: active | under_offer | sold | withdrawn. Only \"active\" is surfaced to users."],
            ["listed_at", "string (ISO 8601)", "Pref",
             "When the listing first appeared. Drives a \"new this week\" badge."],
            ["updated_at", "string (ISO 8601)", "Req",
             "Last modification timestamp. This is the cursor for delta queries. Must bump on any field change."],
        ])
    add_code_block(doc,
        "{\n"
        "    \"listing_id\": \"KAT-2026-001234\",\n"
        "    \"canonical_url\": \"https://katitas.jp/listings/001234\",\n"
        "    \"prefecture\": \"nagano\",\n"
        "    \"city\": \"iiyama\",\n"
        "    \"address_line\": \"1234-5 Example-cho, Iiyama-shi\",\n"
        "    \"lat\": 36.8512,\n"
        "    \"lng\": 138.3661,\n"
        "    \"property_type\": \"detached_house\",\n"
        "    \"condition\": \"livable\",\n"
        "    \"year_built\": 1988,\n"
        "    \"land_area_sqm\": 412.5,\n"
        "    \"building_area_sqm\": 128.0,\n"
        "    \"layout\": \"4LDK\",\n"
        "    \"asking_price_jpy\": 12800000,\n"
        "    \"price_currency\": \"JPY\",\n"
        "    \"listing_agent_name\": \"Katitas Iiyama Office\",\n"
        "    \"hero_image_url\": \"https://img.katitas.jp/001234/hero.jpg\",\n"
        "    \"image_alt_text\": \"4LDK detached house in Iiyama\",\n"
        "    \"status\": \"active\",\n"
        "    \"listed_at\": \"2026-03-14T09:00:00+09:00\",\n"
        "    \"updated_at\": \"2026-04-08T11:42:31+09:00\"\n"
        "}",
        label="JSON payload — one listing")
    add_callout(doc, "Minimum viable subset",
        "listing_id, canonical_url, prefecture, city, property_type, condition, "
        "asking_price_jpy, hero_image_url, status, updated_at. Everything else "
        "improves match quality but is not blocking.")
    doc.add_page_break()

    # --- Page 5: Sync protocol
    add_section_label(doc, "Katitas Technical Spec · 05")
    add_h1(doc, "Sync protocol.")
    add_section_label(doc, "Daily delta · 24-hour cadence")
    add_body(doc, "Japanoma pulls all listings where updated_at > "
                  "last_successful_sync. New records are inserted, existing "
                  "records are upserted by listing_id, records whose status "
                  "flipped off \"active\" are immediately hidden from users.")
    add_section_label(doc, "Weekly full reconciliation · Sunday 03:00 JST")
    add_body(doc, "Japanoma pulls the complete list of active listing IDs in "
                  "the covered regions and diffs against the local index. "
                  "Anything in the local index that is not in Katitas's current "
                  "set is treated as delisted and removed. This catches missed "
                  "deltas, silent deletes, and race conditions.")
    add_section_label(doc, "Freshness SLA")
    add_body(doc, "A listing is never surfaced to users if its updated_at is "
                  "older than 14 days without a verified re-fetch. In practice, "
                  "match cards are at most 24 hours stale, typically under 12.")
    add_section_label(doc, "Error handling")
    add_bullets(doc, [
        "Transient failure. Retry with exponential backoff (1m, 5m, 30m), up "
        "to 3 attempts within the same sync window. Continue serving the last "
        "successful snapshot if all retries fail.",
        "Schema change. Unknown fields are ignored. Missing required fields "
        "cause the affected listing to be skipped and logged. The sync "
        "continues for the remaining listings.",
        "Authentication failure (401/403). Immediate alert to Craefto and to "
        "Katitas's technical contact. No further retries until the credential "
        "is rotated or confirmed.",
        "Rate limit (429). Honour the Retry-After header if present, otherwise "
        "back off for 30 minutes.",
        "Server error (5xx). Same exponential backoff as transient failure; "
        "escalate after three failed windows (72 hours).",
    ])
    add_callout(doc, "Katitas-side obligations for delta to work",
        "(1) updated_at must bump on every change to any user-visible field. "
        "(2) status must flip promptly on under-offer / sold / withdrawn. "
        "(3) listing_id must remain stable for the life of the listing.")
    doc.add_page_break()

    # --- Page 6: Referral URL
    add_section_label(doc, "Katitas Technical Spec · 06")
    add_h1(doc, "Referral URL format.")
    add_body(doc, "Every click from a Japanoma match card to a Katitas listing "
                  "appends three sets of query parameters to the canonical URL. "
                  "No user PII is transmitted.")
    add_code_block(doc,
        "https://katitas.jp/listings/001234\n"
        "    ?ref=japanoma\n"
        "    &utm_source=japanoma\n"
        "    &utm_medium=referral\n"
        "    &utm_campaign=account_match\n"
        "    &japanoma_match_id=m_8f3b2c9d1e4a",
        label="Jump-off URL (line-broken for clarity — emitted as single URL)")
    add_section_label(doc, "Parameter semantics")
    add_schema_table(doc,
        ["Field", "Type", "Req.", "Description"],
        [
            ["ref", "string (const)", "Req",
             "Always \"japanoma\". Primary analytics-agnostic source identifier."],
            ["utm_source", "string", "Req",
             "Always \"japanoma\". Consumed by Katitas's existing analytics tool (GA, Adobe, etc.) without any integration work."],
            ["utm_medium", "string", "Req", "Always \"referral\"."],
            ["utm_campaign", "string", "Req",
             "Identifies the surface: \"account_match\" for curated matches in /account. Future surfaces get different values."],
            ["japanoma_match_id", "string", "Req",
             "Opaque, one-way identifier for the specific match. Katitas cannot use it to look up the user on Japanoma. Safe for outcome feedback."],
        ])
    add_section_label(doc, "Optional outcome feedback webhook")
    add_body(doc, "If Katitas wants to close the loop for attribution, it can "
                  "POST outcome events back to Japanoma using the opaque match "
                  "ID. Entirely optional — not required for launch.")
    add_code_block(doc,
        "POST https://api.japanoma.com/v1/partner/katitas/outcomes\n"
        "Authorization: Bearer <KATITAS_API_KEY>\n"
        "Content-Type: application/json\n"
        "\n"
        "{\n"
        "    \"match_id\": \"m_8f3b2c9d1e4a\",\n"
        "    \"event\": \"contacted\",\n"
        "    \"occurred_at\": \"2026-04-12T14:22:00+09:00\"\n"
        "}\n"
        "\n"
        "Valid event values: clicked | contacted | viewed | offered | sold",
        label="Webhook (optional)")
    doc.add_page_break()

    # --- Page 7: Security
    add_section_label(doc, "Katitas Technical Spec · 07")
    add_h1(doc, "Security, auth, residency.")
    add_section_label(doc, "Authentication · API option")
    add_body(doc, "Bearer token in the Authorization header. Katitas issues "
                  "the key; Japanoma stores it in production environment "
                  "variables only. Never committed to source control, never "
                  "logged. Key rotation supported by re-issue and redeploy.")
    add_section_label(doc, "Authentication · SFTP / S3 option")
    add_bullets(doc, [
        "SFTP: SSH key-pair auth only. Japanoma provides its public key, "
        "Katitas installs it in ~/.ssh/authorized_keys. No password auth.",
        "S3: IAM role assumption or a dedicated IAM user with s3:GetObject on "
        "the feed prefix only. Credentials in Japanoma's production "
        "environment variables.",
    ])
    add_section_label(doc, "Transport")
    add_body(doc, "HTTPS only (TLS 1.2+) for API calls. SFTP only (no FTP). "
                  "S3 with TLS enforced at the bucket policy level.")
    add_section_label(doc, "Network")
    add_body(doc, "Japanoma's outbound calls originate from Vercel's edge "
                  "network. Static IP allowlisting is possible via a proxy if "
                  "Katitas requires it; preference is bearer-token auth without "
                  "allowlisting to avoid the proxy hop.")
    add_section_label(doc, "Data residency")
    add_body(doc, "Japanoma's database is Supabase in the Sydney region "
                  "(ap-southeast-2). Listing records ingested from Katitas are "
                  "stored there. Listings contain no personal information — "
                  "only property descriptions, prices, and media URLs — so "
                  "APPI's cross-border personal-data transfer provisions do "
                  "not apply to the feed data.")
    add_section_label(doc, "Takedown")
    add_body(doc, "Any listing whose status flips to sold, under_offer, or "
                  "withdrawn in the Katitas feed is removed from user-facing "
                  "surfaces within 24 hours (next sync window). Emergency "
                  "takedown of any specific listing within one business day by "
                  "email to Craefto.")
    add_section_label(doc, "Secret management · Japanoma side")
    add_bullets(doc, [
        "All Katitas credentials stored in Vercel environment variables with production scope only.",
        "Credentials never appear in logs, error reports, or Sentry breadcrumbs.",
        "Rotation supported by issuing a new credential and redeploying with the updated env var; the old credential can be revoked immediately after.",
        "Access to production env vars limited to Craefto's technical lead.",
    ])
    doc.add_page_break()

    # --- Page 8: Open items (region bullet removed)
    add_section_label(doc, "Katitas Technical Spec · 08")
    add_h1(doc, "Open items for Katitas to confirm.")
    add_body(doc, "Items Katitas must answer to unblock implementation. "
                  "Grouped by urgency.")
    add_section_label(doc, "Blocking · needed before build starts")
    add_bullets(doc, [
        "Data access method. Which of Option A, B, or C on page 2 can Katitas commit to?",
        "Schema coverage. Can Katitas provide the full schema on pages 3–4, or only the minimum viable subset? Any fields Katitas cannot provide at all?",
        "Stable listing_id. Does Katitas's internal catalogue already have stable, unique IDs per listing that survive updates and re-listings?",
        "updated_at timestamp. Does every listing record have a reliable last-modified timestamp that bumps on any field change?",
        "Technical contact. A named Katitas engineer as the primary technical point for the implementation window and ongoing support.",
    ])
    add_section_label(doc, "Important · needed before go-live")
    add_bullets(doc, [
        "Sandbox or sample payload. A non-production endpoint (Option A) or a sample snapshot file (Option B/C) to build and validate against.",
        "Rate limit. Maximum requests per minute or per hour Japanoma is permitted against the endpoint.",
        "Server maintenance window. So Japanoma can schedule the weekly reconciliation outside Katitas's maintenance slot.",
        "Excluded listing categories. Any segments (restricted, off-market, agent-reserved) that should be filtered out of the feed regardless of the region filter.",
        "Hot-link permission for hero images. Explicit confirmation that Japanoma may hot-link the hero_image_url with \"via Katitas\" attribution.",
    ])
    add_section_label(doc, "Nice to have · can come later")
    add_bullets(doc, [
        "Outcome feedback webhook (page 6) for closed-loop attribution.",
        "Higher-frequency delta sync than daily, if Katitas's infrastructure supports it and Japanoma's match volume justifies it.",
        "Additional fields beyond the schema (e.g. floor plan image, internal quality score) if Katitas considers them safe to share.",
    ])
    add_callout(doc, "One-line summary",
        "Confirm page 2 (access method), confirm pages 3–4 (schema), confirm "
        "page 7 (auth) and name a technical contact. Japanoma is ready to "
        "build against a sandbox within one week of that confirmation.")


def main():
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_margins(doc)
    base_styles(doc)
    build(doc)
    doc.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
